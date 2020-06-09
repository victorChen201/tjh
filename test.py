# -*- coding: utf-8 -*-
import docx
from docxtpl import DocxTemplate
from docxtpl import Document
from docx.shared import Inches

tpl = DocxTemplate('test.docx')
import re
import json
import collections

sd = tpl.new_subdoc()
sd.add_paragraph('Drug Table :')
rows = 3
cols = 4
table = sd.add_table(rows=rows, cols=cols, style=None)

# header
cells = table.rows[0].cells
cells[0].text = "Gene"
cells[1].text = "Drug"
cells[2].text = "Rank"
cells[3].text = "Description"

table.cell(1, 0).text = "ALK"
table.cell(1, 1).text = "GENE1"
table.cell(1, 2).text = "GENE2"
table.cell(1, 3).text = "haha"
#
# #https://python-docx.readthedocs.io/en/latest/api/table.html#docx.table.Table.style
# #merge
table.cell(2, 0).merge(table.cell(2, 3)).text = "One Drug"

context = {
    'mysubdoc': sd,
}

tpl.render(context)
tpl.save('vertical_merge.docx')

# import matplotlib.pyplot as plt
#
# from matplotlib.table import CustomCell
# from matplotlib.widgets import TextBox
#
# def on_pick(event):
#
#     if isinstance(event.artist, CustomCell):
#         cell = event.artist
#         # Doesn't work because cell.get_y() is negative:
#         #text_box_axes = plt.axes([cell.get_x(), cell.get_y(), cell.get_width(), cell.get_height()])
#
#         # This doesn't work either but at least you can see the TextBox on the figure!
#         text_box_axes = plt.axes([cell.get_x(), -cell.get_y(), cell.get_width(), cell.get_height()])
#
#         cell_text = cell.get_text().get_text()
#         TextBox(text_box_axes, '', initial=cell_text)
#         plt.draw()
#
# column_labels = ('Length', 'Width', 'Height', 'Sold?')
# row_labels = ['Ferrari', 'Porsche']
# data = [[2.2, 1.6, 1.2, True],
#         [2.1, 1.5, 1.4, False]]
# table = plt.table(cellText=data, colLabels=column_labels, rowLabels=row_labels, cellLoc='center', loc='best')
# text_box = None
#
# celld = table.get_celld()
# for key in celld.keys():
#     # Each key is a tuple of the form (row, column).
#     # Column headings are in row 0. Row headings are in column -1.
#     # So the first item of data in the table is actually at (1, 0).
#     if key[0] > 0 and key[1] > -1:
#         cell = celld[key]
#         cell.set_picker(True)
#         cell._width = cell._width/2
#         # print(dir(cell))
#
# canvas = plt.gcf().canvas
# canvas.mpl_connect('pick_event', on_pick)
# plt.axis('off')
#
# plt.show()
# import matplotlib.pyplot as plt
# import numpy as np
# import pandas
#
# from matplotlib.table import Table
#
# def main():
#  data = pandas.DataFrame(np.random.random((12,8)),
#  columns=['A','B','C','D','E','F','G','H'])
#  checkerboard_table(data)
#  plt.show()
#
# def checkerboard_table(data, fmt='{:.2f}', bkg_colors=['yellow', 'white']):
#  fig, ax = plt.subplots()
#  ax.set_axis_off()
#  tb = Table(ax, bbox=[0,0,1,1])
#
#  nrows, ncols = data.shape
#  width, height = 1.0/ncols, 1.0/nrows
#
#  # Add cells
#  for (i,j), val in np.ndenumerate(data):
#      # Index either the first or second item of bkg_colors based on
#      # a checker board pattern
#      idx = [j % 2, (j + 1) % 2][i % 2]
#      color = bkg_colors[idx]
#
#      tb.add_cell(i, j, width, height, text=fmt.format(val),
#      loc='center', facecolor=color)
#
#  # Row Labels...
#  for i, label in enumerate(data.index):
#      tb.add_cell(i, -1, width, height, text=label, loc='right',
#      edgecolor='none', facecolor='none')
#      # Column Labels...
#  for j, label in enumerate(data.columns):
#      tb.add_cell(-1, j, width, height/2, text=label, loc='center',
#      edgecolor='none', facecolor='none')
#      ax.add_table(tb)
#  return fig
#
# if __name__ == '__main__':
#  main()